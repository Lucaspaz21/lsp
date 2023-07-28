from lsp import app
from flask import render_template, request, url_for, redirect, flash
from lsp.forms import FormLogin, Upar
from werkzeug.utils import secure_filename
import os
import pandas as pd
import matplotlib.pyplot as plt
import matplotlib
from decimal import Decimal, InvalidOperation
import smtplib
import email.message
from datetime import date, datetime
import docx
from docx2pdf import convert
matplotlib.use('Agg')

user, passw = 'lsp', 'lsp'

def moeda(n=0, moeda='R$'):
    return f'{moeda}{n:.2f}'.replace('.', ',')

def relatorio_mensal():
    atual = datetime.now()
    dia = atual.strftime("%d/%m/%Y")
    try:
        caminho = r'lsp/static/files/Vendas_3.1.xlsm'
        base_rela_mensal = pd.read_excel(caminho, sheet_name='Dados')
    except:
        erro_rela_mensal = 'Não foi possivel obter a base'
        print(erro_rela_mensal)

    try:
        # selecionar dados
        vendamensal = base_rela_mensal.iloc[2, 1]
        saidamensal = base_rela_mensal.iloc[3, 1]
        saldo = vendamensal - saidamensal
        combustivel = base_rela_mensal.iloc[4, 1]
        loja = base_rela_mensal.iloc[5, 1]
        basico = base_rela_mensal.iloc[6, 1]
        sistema = base_rela_mensal.iloc[8, 1]
        outros = base_rela_mensal.iloc[8, 1]
        contasfuturas = base_rela_mensal.iloc[1, 1]

        debito = base_rela_mensal.iloc[9, 1]
        credito = base_rela_mensal.iloc[10, 1]
        dinheiro = base_rela_mensal.iloc[11, 1]
        parcelado = base_rela_mensal.iloc[12, 1]
        pix = base_rela_mensal.iloc[13, 1]
        cheque = base_rela_mensal.iloc[14, 1]

        # dados de porcentagem
        dados_loja = float((loja / saidamensal) * 100)
        dados_basico = (basico / saidamensal) * 100
        dados_sistema = (sistema / saidamensal) * 100
        dados_outros = (outros / saidamensal) * 100
        dados_combustivel = (combustivel / saidamensal) * 100

        dados_debito = (debito / vendamensal) * 100
        dados_credito = (credito / vendamensal) * 100
        dados_pix = (pix / vendamensal) * 100
        dados_cheque = (cheque / vendamensal) * 100
        dados_dinheiro = (dinheiro / vendamensal) * 100
        dados_parcelado = (parcelado / vendamensal) * 100
    except:
        erro_importar_infor_rela = 'ERRO ao importar dados para Relatório Mensal!'
        print(erro_importar_infor_rela)


    # criando arquivo Word
    documento = docx.Document()
    documento.add_heading(f"Relatório Mensal - {dia}", 0)
    documento.add_paragraph(f"Relatório gerado automaticamente com base nos dados lançados no sistema.")
    documento.add_heading("Dados gerais do mês.", 3)
    documento.add_paragraph(f"Venda mensal no valor de {moeda(vendamensal)}.")
    documento.add_paragraph(f"Saida mensal no valor de {moeda(saidamensal)}.")
    documento.add_paragraph(f"Saldo mensal no valor de {moeda(saldo)}.")
    documento.add_paragraph(f"Contas futuras em aberto no valor de {moeda(contasfuturas)}.")
    documento.add_paragraph("------------------------------------------------------------------------")
    # detalhes despesas do mes
    documento.add_heading("Detalhes de despesas do mês.", 3)
    documento.add_paragraph(f"Gasto em combustível: {moeda(combustivel)}.")
    documento.add_paragraph(f"Gasto de sistema da loja (salarios, impostos, aluguel etc): {moeda(sistema)}.")
    documento.add_paragraph(f"Outros gastos para fins não planejados: {moeda(outros)}.")
    documento.add_paragraph(f"Gastos com produtos básicos (cimento, areia, pedra): {moeda(basico)}.")
    documento.add_paragraph(f"Gastos com produtos não básicos: {moeda(loja)}.")
    # estatísticas do mes
    documento.add_heading("Estatísticas de despesas.", 3)
    documento.add_paragraph(f"Produtos considerados de loja obtiveram um valor de"
                            f" {dados_loja:.2f}% da Saida Mensal, enquanto que Produtos"
                            f" considerados Básicos ficaram em {dados_basico:.2f}% da Saída Mensal.")
    documento.add_paragraph(
        f"Os gastos que envolvem salários, aluguel e impostos tiveram porcentagem de {dados_sistema:.2f}%, "
        f" gastos não planejados ficaram em {dados_outros:.2f}%.")
    documento.add_paragraph(f"O gasto em combustível é referente a {dados_combustivel:.2f}% da Saída Mensal.")
    documento.add_paragraph("------------------------------------------------------------------------")

    # detalhe sobre formas de entrada
    documento.add_heading("Estatísticas de entradas.", 3)
    documento.add_paragraph(f"Entrada em Dinheiro: {moeda(dinheiro)}, equivalente a {dados_dinheiro:.2f}%.")
    documento.add_paragraph(f"Entrada em Débito: {moeda(debito)}, equivalente a {dados_debito:.2f}%. ")
    documento.add_paragraph(f"PIX: {moeda(pix)}, equivalente a {dados_pix:.2f}%.")
    documento.add_paragraph(
        f"Entrada em Crédito: R${moeda(credito)}, eqivalente a {dados_credito:.2f}%, Crédito Parcelado: {moeda(parcelado)},"
        f" equivalente a {dados_parcelado:.2f}%, Cheque: {moeda(dados_cheque)}.")
    documento.add_paragraph(f"Arquivo gerado em {dia}.")
    # salvar arquivo
    dia1 = atual.strftime("%d-%m-%Y")
    nome = f'Relatório Mensal - {dia1}.docx'
    documento.save(f"lsp/static/files/relatorios/{nome}")

def tratardados():
    caminho = r'lsp/static/files/Vendas_3.1.xlsm'
    arquivo = pd.read_excel(caminho, sheet_name='Dados')
    venda_mensal = arquivo.iloc[2, 1]
    saida_mensal = arquivo.iloc[3, 1]
    saida_combustivel = arquivo.iloc[4, 1]
    saida_loja = arquivo.iloc[5, 1]
    saida_basico = arquivo.iloc[6, 1]
    saida_sistema = arquivo.iloc[7, 1]
    saida_outros = arquivo.iloc[8, 1]
    entrada_debito = arquivo.iloc[9, 1]
    entrada_credito = arquivo.iloc[10, 1]
    entrada_dinheiro = arquivo.iloc[11, 1]
    entrada_credido_parcelado = arquivo.iloc[12, 1]
    entrada_pix = arquivo.iloc[13, 1]
    entrada_cheque = arquivo.iloc[14, 1]

    def grafico_geral():
        # GRAFICO GERAL
        categorias = ['Venda Mensal', 'Saída Mensal']
        valores = [venda_mensal, saida_mensal]

        # Criar gráfico de pizza
        fig1, ax1 = plt.subplots()
        patches, texts, autotexts = ax1.pie(valores, labels=categorias, autopct='%1.1f%%')

        # Adicionar anotações dos valores no gráfico
        for i, texto in enumerate(autotexts):
            texto.set_text(f'R${valores[i]:.2f}')

        ax1.set_title('Gráfico de Rendimento')

        # Salvar Gráfico
        fig1.savefig('lsp/static/images/grafico.png')

        plt.close(fig1)

    def grafico_entradas():
        # Dados de Entradas
        categorias_entrada = ['Entrada Débito', 'Entrada Dinheiro', 'Entrada Crédito', 'Entrada Credito Parcelado',
                              'Entrada PIX', 'Entrada Cheque']
        valores_entrada = [entrada_debito, entrada_dinheiro, entrada_credito, entrada_credido_parcelado, entrada_pix,
                           entrada_cheque]

        # Criar gráfico de pizza
        fig2, ax2 = plt.subplots()
        ax2.pie(valores_entrada, labels=categorias_entrada, autopct='%1.1f%%')
        ax2.set_title('Gráfico de Entradas')

        # Salvar Gráfico
        fig2.savefig('lsp/static/images/grafico_entradas.png')

        plt.close(fig2)

    def grafico_saidas():
        # Criar gráfico de saída

        # Dados de Saída
        categorias_saida = ['Saída Combustível', 'Saída Loja', 'Saída Básico', 'Saída Sistema', 'Saída Outros']
        valores_saida = [saida_combustivel, saida_loja, saida_basico, saida_sistema, saida_outros]

        fig3, ax3 = plt.subplots()
        ax3.pie(valores_saida, labels=categorias_saida, autopct='%1.1f%%')
        ax3.set_title('Gráfico de Saídas')

        # Salvar Gráfico
        fig3.savefig('lsp/static/images/grafico_saidas.png')

        # Fechar as figuras para liberar memória


        plt.close(fig3)

    grafico_geral()
    grafico_entradas()
    grafico_saidas()

def criartuts():
    caminho = r'lsp/static/files/Vendas_3.1.xlsm'
    arquivo = pd.read_excel(caminho, sheet_name='Tuts')
    col_códigos = arquivo.iloc[:, 0]
    col_produtos = arquivo.iloc[:, 1]
    col_custo = arquivo.iloc[:, 3]
    col_venda = arquivo.iloc[:, 4]
    col_desconto = arquivo.iloc[:, 9]
    tuts = pd.merge(col_códigos, col_produtos, left_index=True, right_index=True)
    tuts = pd.merge(tuts, col_custo, left_index=True, right_index=True)
    tuts = pd.merge(tuts, col_venda, left_index=True, right_index=True)
    tuts = pd.merge(tuts, col_desconto, left_index=True, right_index=True)

    tuts.to_csv(r'lsp/static/files/Tuts.csv', sep=';', index=False)

def converter_valor(valor_str):
    try:
        valor_decimal = Decimal(valor_str.replace(',', '.'))
        return valor_decimal
    except InvalidOperation:
        return None

def taxas(valor):
    caminho = r'lsp/static/files/Vendas_3.1.xlsm'
    arquivo = pd.read_excel(caminho, sheet_name='Dados')
    ton_debito = Decimal(arquivo.iloc[1, 4])
    ton_credito = Decimal(arquivo.iloc[1, 5])
    ton_credito_2x = Decimal(arquivo.iloc[1, 6])
    ton_credito_3x = Decimal(arquivo.iloc[1, 7])
    ton_credito_4x = Decimal(arquivo.iloc[1, 8])
    ton_credito_5x = Decimal(arquivo.iloc[1, 9])
    ton_credito_6x = Decimal(arquivo.iloc[1, 10])
    ton_credito_7x = Decimal(arquivo.iloc[1, 11])
    ton_credito_8x = Decimal(arquivo.iloc[1, 12])
    ton_credito_9x = Decimal(arquivo.iloc[1, 13])
    ton_credito_10x = Decimal(arquivo.iloc[1, 14])
    ton_credito_11x = Decimal(arquivo.iloc[1, 15])
    ton_credito_12x = Decimal(arquivo.iloc[1, 16])

    cielo_debito = Decimal(arquivo.iloc[2, 4])
    cielo_credito = Decimal(arquivo.iloc[2, 5])
    cielo_credito_2x = Decimal(arquivo.iloc[2, 6])
    cielo_credito_3x = Decimal(arquivo.iloc[2, 7])
    cielo_credito_4x = Decimal(arquivo.iloc[2, 8])
    cielo_credito_5x = Decimal(arquivo.iloc[2, 9])
    cielo_credito_6x = Decimal(arquivo.iloc[2, 10])
    cielo_credito_7x = Decimal(arquivo.iloc[2, 11])
    cielo_credito_8x = Decimal(arquivo.iloc[2, 12])
    cielo_credito_9x = Decimal(arquivo.iloc[2, 13])
    cielo_credito_10x = Decimal(arquivo.iloc[2, 14])
    cielo_credito_11x = Decimal(arquivo.iloc[2, 15])
    cielo_credito_12x = Decimal(arquivo.iloc[2, 16])



    # Calcular pagamentos
    ton_debito_pagamento = f'R${valor - (valor * ton_debito / 100):.2f}'
    ton_credito_pagamento = f'R${valor - (valor * ton_credito / 100):.2f}'
    ton_credito_2x_pagamento = f'R${valor - (valor * ton_credito_2x / 100):.2f}'
    ton_credito_3x_pagamento = f'R${valor - (valor * ton_credito_3x / 100):.2f}'
    ton_credito_4x_pagamento = f'R${valor - (valor * ton_credito_4x / 100):.2f}'
    ton_credito_5x_pagamento = f'R${valor - (valor * ton_credito_5x / 100):.2f}'
    ton_credito_6x_pagamento = f'R${valor - (valor * ton_credito_6x / 100):.2f}'
    ton_credito_7x_pagamento = f'R${valor - (valor * ton_credito_7x / 100):.2f}'
    ton_credito_8x_pagamento = f'R${valor - (valor * ton_credito_8x / 100):.2f}'
    ton_credito_9x_pagamento = f'R${valor - (valor * ton_credito_9x / 100):.2f}'
    ton_credito_10x_pagamento = f'R${valor - (valor * ton_credito_10x / 100):.2f}'
    ton_credito_11x_pagamento = f'R${valor - (valor * ton_credito_11x / 100):.2f}'
    ton_credito_12x_pagamento = f'R${valor - (valor * ton_credito_12x / 100):.2f}'

    cielo_debito_pagamento = f'R${valor - (valor * cielo_debito / 100):.2f}'
    cielo_credito_pagamento = f'R${valor - (valor * cielo_credito / 100):.2f}'
    cielo_credito_2x_pagamento = f'R${valor - (valor * cielo_credito_2x / 100):.2f}'
    cielo_credito_3x_pagamento = f'R${valor - (valor * cielo_credito_3x / 100):.2f}'
    cielo_credito_4x_pagamento = f'R${valor - (valor * cielo_credito_4x / 100):.2f}'
    cielo_credito_5x_pagamento = f'R${valor - (valor * cielo_credito_5x / 100):.2f}'
    cielo_credito_6x_pagamento = f'R${valor - (valor * cielo_credito_6x / 100):.2f}'
    cielo_credito_7x_pagamento = f'R${valor - (valor * cielo_credito_7x / 100):.2f}'
    cielo_credito_8x_pagamento = f'R${valor - (valor * cielo_credito_8x / 100):.2f}'
    cielo_credito_9x_pagamento = f'R${valor - (valor * cielo_credito_9x / 100):.2f}'
    cielo_credito_10x_pagamento = f'R${valor - (valor * cielo_credito_10x / 100):.2f}'
    cielo_credito_11x_pagamento = f'R${valor - (valor * cielo_credito_11x / 100):.2f}'
    cielo_credito_12x_pagamento = f'R${valor - (valor * cielo_credito_12x / 100):.2f}'

    return ton_debito_pagamento,  ton_credito_pagamento, ton_credito_2x_pagamento, ton_credito_3x_pagamento, ton_credito_4x_pagamento,\
           ton_credito_5x_pagamento, ton_credito_6x_pagamento, ton_credito_7x_pagamento, ton_credito_8x_pagamento, ton_credito_9x_pagamento,\
           ton_credito_10x_pagamento, ton_credito_11x_pagamento, ton_credito_12x_pagamento, cielo_debito_pagamento,  cielo_credito_pagamento, cielo_credito_2x_pagamento, cielo_credito_3x_pagamento, cielo_credito_4x_pagamento,\
           cielo_credito_5x_pagamento, cielo_credito_6x_pagamento, cielo_credito_7x_pagamento, cielo_credito_8x_pagamento, cielo_credito_9x_pagamento,\
           cielo_credito_10x_pagamento, cielo_credito_11x_pagamento, cielo_credito_12x_pagamento

def relatorios_semanal():
    caminho = r'lsp/static/files/Vendas_3.1.xlsm'
    try:
        tabela = pd.read_excel(caminho, sheet_name='Dados')
        # tabela = pd.read_excel(r"C:\L da S Paz\Loja\Venda diarias v3.1/Vendas_3.1.xlsm", sheet_name='Dados')
    except:
        print('Erro ao abrir tabela')

    try:
        vendamensal = tabela.iloc[2, 1]
        saidamensal = tabela.iloc[3, 1]
        contasfuturas = tabela.iloc[1, 1]
        saldoatual = vendamensal - saidamensal
        dia = f'{date.today().day}/{date.today().month}/{date.today().year}'
    except:
        print('Falha ao importar dados para Relatório Semanal!')

    def relatorio_semanal():
        try:
            def enviar_email():  # gilberto
                corpo_email = f'''<h4>Prezados,</h4>

                <p>Segue o Relatório de Vendas Mensal do dia <b>{dia}</b>:</p>
                <br>
                <p>Venda mensal até o momento: <b>{moeda(vendamensal)}</b>.</p>
                <p>Saída mensal até o momento: <b>{moeda(saidamensal)}</b>.</p>
                <p>Saldo mensal até o momento: <b>{moeda(saldoatual)}</b>.</p>
                <p>Contas em aberto até o momento: <b>{moeda(contasfuturas)}</b>.</p>
                <br>
                <p><b>Att.,</b></p>
                <p><b>Lucas Paz</b></p>
                '''

                msg = email.message.Message()
                msg['Subject'] = f"Venda mensal - {dia}"
                msg['From'] = 'lucas.paz12@gmail.com'
                msg['To'] = 'gilbertomazetti@gmail.com'
                password = 'hhevbfduxzqqyetd'
                msg.add_header('Content-Type', 'text/html')
                msg.set_payload(corpo_email)

                s = smtplib.SMTP('smtp.gmail.com: 587')
                s.starttls()
                # Login Credentials for sending the mail
                s.login(msg['From'], password)
                s.sendmail(msg['From'], [msg['To']], msg.as_string().encode('utf-8'))
                print('Email enviado para -> Gilberto Mazetti')
            enviar_email()  # gilberto

            def enviar_email1():  # loja
                corpo_email = f'''<h4>Prezados,</h4>

                <p>Segue o Relatório de Vendas Mensal do dia <b>{dia}</b>:</p>
                <br>
                <p>Venda mensal até o momento: <b>{moeda(vendamensal)}</b>.</p>
                <p>Saída mensal até o momento: <b>{moeda(saidamensal)}</b>.</p>
                <p>Saldo mensal até o momento: <b>{moeda(saldoatual)}</b>.</p>
                <p>Contas em aberto até o momento: <b>{moeda(contasfuturas)}</b>.</p>
                <br>
                <p><b>Att.,</b></p>
                <p><b>Lucas Paz</b></p>
                '''

                msg = email.message.Message()
                msg['Subject'] = f"Venda mensal - {dia}"
                msg['From'] = 'lucas.paz12@gmail.com'
                msg['To'] = 'casadeconstrucaocarmo@gmail.com'
                password = 'hhevbfduxzqqyetd'
                msg.add_header('Content-Type', 'text/html')
                msg.set_payload(corpo_email)

                s = smtplib.SMTP('smtp.gmail.com: 587')
                s.starttls()
                # Login Credentials for sending the mail
                s.login(msg['From'], password)
                s.sendmail(msg['From'], [msg['To']], msg.as_string().encode('utf-8'))
                print('Email enviado para -> L da S Paz Materiais de Construção')

                # enviar_email1()  # loja
            enviar_email1()

            def enviar_email2():  # marcia
                corpo_email = f'''<h4>Prezados,</h4>

                <p>Segue o Relatório de Vendas Mensal do dia <b>{dia}</b>:</p>
                <br>
                <p>Venda mensal até o momento: <b>{moeda(vendamensal)}</b>.</p>
                <p>Saída mensal até o momento: <b>{moeda(saidamensal)}</b>.</p>
                <p>Saldo mensal até o momento: <b>{moeda(saldoatual)}</b>.</p>
                <p>Contas em aberto até o momento: <b>{moeda(contasfuturas)}</b>.</p>
                <br>
                <p><b>Att.,</b></p>
                <p><b>Lucas Paz</b></p>
                '''

                msg = email.message.Message()
                msg['Subject'] = f"Venda mensal - {dia}"
                msg['From'] = 'lucas.paz12@gmail.com'
                msg['To'] = 'cscleidemarcia@gmail.com'
                password = 'hhevbfduxzqqyetd'
                msg.add_header('Content-Type', 'text/html')
                msg.set_payload(corpo_email)

                s = smtplib.SMTP('smtp.gmail.com: 587')
                s.starttls()
                # Login Credentials for sending the mail
                s.login(msg['From'], password)
                s.sendmail(msg['From'], [msg['To']], msg.as_string().encode('utf-8'))
                print('Email enviado para -> Cleide Marcia')
            enviar_email2()  # marcia

            def enviar_email3():  # lucas
                corpo_email = f'''<h4>Prezados,</h4>

                <p>Segue o Relatório de Vendas Mensal do dia <b>{dia}</b>:</p>
                <br>
                <p>Venda mensal até o momento: <b>{moeda(vendamensal)}</b>.</p>
                <p>Saída mensal até o momento: <b>{moeda(saidamensal)}</b>.</p>
                <p>Saldo mensal até o momento: <b>{moeda(saldoatual)}</b>.</p>
                <p>Contas em aberto até o momento: <b>{moeda(contasfuturas)}</b>.</p>
                <br>
                <p><b>Att.,</b></p>
                <p><b>Lucas Paz</b></p>
                '''

                msg = email.message.Message()
                msg['Subject'] = f"Venda mensal - {dia}"
                msg['From'] = 'lucas.paz12@gmail.com'
                msg['To'] = 'lucas.paz12@gmail.com'
                password = 'hhevbfduxzqqyetd'
                msg.add_header('Content-Type', 'text/html')
                msg.set_payload(corpo_email)

                s = smtplib.SMTP('smtp.gmail.com: 587')
                s.starttls()
                # Login Credentials for sending the mail
                s.login(msg['From'], password)
                s.sendmail(msg['From'], [msg['To']], msg.as_string().encode('utf-8'))
                print('Email enviado para -> Lucas Paz')
            enviar_email3()  # lucas
        except:
            print('Erro ao enviar email!')

    relatorio_semanal()


@app.route('/9Y5h2--home-D6f7R1A')
def home():
    tratardados()
    return render_template('home.html')

@app.route('/login', methods=['GET', 'POST'])
def login():
    form_login = FormLogin()
    if form_login.validate_on_submit():
        if form_login.username.data == user and form_login.senha.data == passw:
            return redirect(url_for('home'))
        else:
            flash(f'Falha no Login. E-mail ou Senha Incorretos', 'alert-danger')
            return redirect(url_for('login'))
    return render_template('login.html', form_login = form_login)

@app.route('/4aS8j--admin-P2kL5tR', methods=['GET', 'POST'])
def admin():
    form = Upar()
    if form.validate_on_submit():
        try:
            file = form.arquivo.data
            file.save(os.path.join(os.path.abspath(os.path.dirname(__file__)), app.config['UPLOAD_FOLDER'], secure_filename(file.filename)))
            return redirect(url_for('login'))
        except:
            return 'Erro ao fazer upload'
    return render_template('admin.html', form_upar=form)

@app.route('/6bF9zG3--graficos--cV7mX')
def graficos():
    tratardados()
    return render_template('graficos.html')

@app.route('/1qW7hE--produtos--5dY2oT')
def produtos():
    criartuts()
    tuts = pd.read_csv(r'lsp/static/files/Tuts.csv', sep=';')
    data = {
        'Código': tuts.iloc[:, 0],
        'Produto': tuts.iloc[:, 1],
        'Custo': tuts.iloc[:, 2],
        'Venda': tuts.iloc[:, 3],
        'Venda c/ Desconto': tuts.iloc[:, 4],
    }

    df = pd.DataFrame(data)

    produto = request.args.get('Produto')
    codigo = request.args.get('Código')

    if produto:
        filtro_produto = df['Produto'].str.contains(produto, case=False)
    else:
        filtro_produto = pd.Series(True, index=df.index)  # Retorna tudo se nenhum produto for especificado

    if codigo:
        filtro_codigo = df['Código'].astype(str).str.contains(codigo, case=False, regex=True)
    else:
        filtro_codigo = pd.Series(True, index=df.index)  # Retorna tudo se nenhum código for especificado

    filtro = filtro_produto & filtro_codigo
    obj = df[filtro].to_dict('split')

    return render_template('produtos.html', obj=obj, produto=produto, codigo=codigo)

@app.route('/taxas', methods=['GET', 'POST'])
def caltaxas():
    if request.method == 'POST':
        valor_str = request.form['valor']
        valor_decimal = converter_valor(valor_str)
        if valor_decimal is not None:
            pagamentos = taxas(valor_decimal)
        return render_template('taxas.html', ton_debito_pagamento=pagamentos[0],
                               ton_credito_pagamento=pagamentos[1],
                               ton_credito_2x_pagamento=pagamentos[2],
                               ton_credito_3x_pagamento=pagamentos[3],
                               ton_credito_4x_pagamento=pagamentos[4],
                               ton_credito_5x_pagamento=pagamentos[5],
                               ton_credito_6x_pagamento=pagamentos[6],
                               ton_credito_7x_pagamento=pagamentos[7],
                               ton_credito_8x_pagamento=pagamentos[8],
                               ton_credito_9x_pagamento=pagamentos[9],
                               ton_credito_10x_pagamento=pagamentos[10],
                               ton_credito_11x_pagamento=pagamentos[11],
                               ton_credito_12x_pagamento=pagamentos[12],
                               cielo_debito_pagamento=pagamentos[13],
                               cielo_credito_pagamento=pagamentos[14],
                               cielo_credito_2x_pagamento=pagamentos[15],
                               cielo_credito_3x_pagamento=pagamentos[16],
                               cielo_credito_4x_pagamento=pagamentos[17],
                               cielo_credito_5x_pagamento=pagamentos[18],
                               cielo_credito_6x_pagamento=pagamentos[19],
                               cielo_credito_7x_pagamento=pagamentos[20],
                               cielo_credito_8x_pagamento=pagamentos[21],
                               cielo_credito_9x_pagamento=pagamentos[22],
                               cielo_credito_10x_pagamento=pagamentos[23],
                               cielo_credito_11x_pagamento=pagamentos[24],
                               cielo_credito_12x_pagamento=pagamentos[25])

    return render_template('taxas.html')

@app.route('/arquivos')
def mostrar_arquivos():
    pasta = 'lsp/static/files/relatorios'  
    arquivos = os.listdir(pasta)
    caminhos_arquivos = [os.path.join(pasta, arquivo) for arquivo in arquivos]
    print(arquivos)
    return render_template('relatorio_mensal.html', arquivos=arquivos, caminhos_arquivos=caminhos_arquivos)

@app.route('/relatorios', methods=['GET', 'POST'])
def relatorios():
    button_pressed = None
    if request.method == 'POST':
        button_pressed = request.form.get('button')

    if button_pressed == 'relatorio_mensal':
        relatorio_mensal()
        # return resultado
        return redirect(url_for('home'))

    elif button_pressed == 'relatorio_semanal':
        relatorios_semanal()
        return redirect(url_for('home'))
        # return resultado
    return render_template('relatorios.html')
